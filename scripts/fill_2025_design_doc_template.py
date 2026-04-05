from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = Path("/Users/saksk/Downloads/01-3 软件应用与开发类作品设计和开发文档模板（2025版）.docx")
OUTPUT_DIR = ROOT / "output" / "doc" / "2025_设计和开发文档模板填写版"
OUTPUT_DOCX = OUTPUT_DIR / "01-3 软件应用与开发类作品设计和开发文档模板（2025版）-已填写-Sak-AI答题助手.docx"
DOWNLOAD_DOCX = Path("/Users/saksk/Downloads/01-3 软件应用与开发类作品设计和开发文档模板（2025版）-已填写-Sak-AI答题助手.docx")
OUTPUT_PDF = OUTPUT_DIR / "01-3 软件应用与开发类作品设计和开发文档模板（2025版）-已填写-Sak-AI答题助手.pdf"

ARCH_IMG = ROOT / "output" / "doc" / "1软件应用与开发_1Web应用与开发_Sak-AI答题助手" / "02素材与源码" / "插图" / "系统架构图.png"
FLOW_IMG = ROOT / "output" / "doc" / "1软件应用与开发_1Web应用与开发_Sak-AI答题助手" / "02素材与源码" / "插图" / "核心流程图.png"
SCREENSHOT_DIR = ROOT / ".submission_assets" / "screenshots"

PROJECT_NUMBER = ""
PROJECT_NAME = "Sak-AI答题助手"
AUTHORS = "王为硕、队员B"
VERSION = "V1.0"
DATE_TEXT = "2026年4月6日"
PUBLIC_URL = "https://saksk.top"
LOCAL_URL = "http://127.0.0.1:8000"

NEEDS = [
    "Sak-AI答题助手面向大学生备考、课程练习与题库运营场景，重点解决题库分散、错题复盘弱、考试训练不足和学习数据不可视等问题。传统刷题工具往往偏重单点练习，公共题库与个人题库相互割裂，导致用户难以沉淀个人学习资源，也不利于教师或运营侧开展内容管理。本作品以 Web 应用为主体，通过统一的题库管理、练习、考试、数据分析、论坛互动和后台管理能力，为学习者和题库维护者提供完整闭环。",
    "目标用户主要包括三类：一是需要高频刷题和复盘的大学生与备考学习者；二是需要组织题库、审核内容和管理用户的教师或运营人员；三是希望随时随地访问同一套题库语义的移动端用户。对标常见在线刷题平台和通用题库系统，本作品强调工程化部署、共享数据语义以及学习闭环的一体化表达。",
]

COMPETITOR_ROWS = [
    ["对比维度", "常见刷题平台", PROJECT_NAME],
    ["部署方式", "多为独立前端 + API 服务", "单仓全栈，Docker 统一部署"],
    ["题库结构", "偏公共题库或单一业务线", "公共题库与个人题库双线并存"],
    ["学习闭环", "练习与复盘分散", "练习、错题、收藏、考试、数据中心一体化"],
    ["运营能力", "社区与后台常需外挂系统", "内建论坛与后台管理模块"],
    ["移动场景", "通常需独立 App 或 H5 适配", "微信小程序复用同一后端语义"],
]

OVERVIEW = [
    "系统整体采用 Flask 单仓全栈架构：Web 浏览器与微信小程序共同访问 Flask 提供的页面与 API 服务，后端连接 PostgreSQL 存储题库与学习数据，使用 Redis 承担缓存、限流与异步任务队列，RQ Worker 负责后台任务处理；本地开发通过 compose.dev.yml 启动完整环境。",
    "核心模块按照业务边界拆分为认证、主页面、题库练习、考试、用户、聊天、通知、弹窗、编程、个人题库、论坛与后台管理等 12 个 Blueprint 模块，兼顾功能内聚与后期扩展。",
    "模块划分如下：",
]

OVERVIEW_BULLETS = [
    "用户认证：登录、会话、JWT、角色权限。",
    "公共题库：题库广场、题库详情、题目浏览。",
    "个人题库：自建题库、公开分享、加入题库与管理。",
    "刷题与错题收藏：答题记录、错题、收藏、进度同步。",
    "模拟考试：考试创建、选择题库、提交与结果查看。",
    "数据中心：正确率、覆盖率、趋势图、热力图与能力画像。",
    "论坛社区：帖子、评论、版块与互动。",
    "后台管理：题目、用户、科目和系统统计管理。",
]

DETAIL = [
    "用户在首页进入系统后，可以从题库广场浏览公共题库，也可以在“我的题库”中访问已创建或已加入的题库资源；随后进入练习与考试流程，在答题后由数据中心进行统计与复盘，形成连续的学习闭环。",
    "数据库设计以用户、题库、题目、学习记录、考试记录和论坛内容为核心实体，通过统一的数据语义支撑 Web 与小程序共用同一套后端逻辑。",
    "关键技术点包括：",
]

DATABASE_ROWS = [
    ["实体/对象", "作用说明"],
    ["用户（User）", "保存账户、角色、资料与登录态信息，支撑 Session/JWT 双认证。"],
    ["题库（公共/个人）", "区分系统题库、用户公开题库与个人题库，承载不同使用场景。"],
    ["题目（Question）", "保存题型、题干、选项、答案与解析等核心数据。"],
    ["学习记录", "记录答题结果、用户答案、做题时间、正确率与复盘数据。"],
    ["考试记录", "记录模拟考试的创建、提交、得分与结果统计。"],
    ["论坛内容", "记录帖子、评论、点赞与互动数据，形成学习社区。"],
]

DETAIL_NUMBERED = [
    "采用 Flask 应用工厂 + Blueprint 组织方式，便于按模块扩展。",
    "兼容 Session 与 JWT 双认证模式，分别服务 Web 与 API/小程序。",
    "接口返回逐步统一为 status/code/data/message 信封结构，降低前后端兼容成本。",
    "基于 Docker 的本地开发模式统一了 Web、Worker、PostgreSQL 与 Redis 运行环境。",
    "通过数据中心可视化组件将练习、考试、复盘转化为可感知的学习反馈。",
    "AI 题目解析能力通过兼容接口接入，用于增强题目讲解与学习辅助体验。",
]

SCREENSHOTS = [
    ("01-首页.png", "/hub", "登录后的首页汇总题库入口、继续学习、签到与学习统计，是用户进入系统后的主控面板。"),
    ("02-题库广场.png", "/public/banks", "题库广场按系统题库与用户公开题库统一呈现，支持浏览、筛选与加入题库。"),
    ("03-题库名片详情.png", "/public/banks/card/user/46", "题库名片详情页展示题库简介、参与人数、加入方式与继续练习入口。"),
    ("04-我的题库.png", "/user/banks", "我的题库聚合用户创建与加入的题库资源，方便继续练习、整理和管理。"),
    ("05-数据中心.png", "/data", "数据中心通过趋势图、热力图和能力画像等可视化组件反馈学习状态。"),
    ("06-考试选择页.png", "/exams/select", "考试中心支持按公共题库或个人题库选择考试来源，进入模拟考试流程。"),
    ("07-论坛首页.png", "/forum", "论坛模块支持按版块浏览帖子、查看热帖并发起交流，强化学习社区互动。"),
    ("08-后台仪表盘.png", "/admin/dashboard", "后台仪表盘用于查看题目、科目、用户等系统级统计信息，支撑平台运营。"),
    ("09-首页-移动端.png", "/hub（移动端）", "移动端适配截图展示 Web 页面在窄屏设备下的响应式布局与触控友好交互。"),
]

TEST_ROWS = [
    ["测试项", "验证方法", "结果"],
    ["容器服务运行", "执行 docker compose --env-file .env -f compose.dev.yml ps", "web/postgres/redis/worker 运行正常"],
    ["基础接口", f"访问 {LOCAL_URL}/api/ping", "返回 status=success"],
    ["深度健康检查", f"访问 {LOCAL_URL}/api/ping?deep=1", "返回 db=true、redis=true"],
    ["首页与导航", "浏览 /hub 并检查题库、考试、论坛入口", "页面可访问，导航正确"],
    ["题库广场", "浏览 /public/banks 与题库详情页", "列表、详情与题库信息展示正常"],
    ["个人题库", "浏览 /user/banks", "个人题库聚合展示正常"],
    ["数据中心", "浏览 /data", "图表容器与统计指标正常加载"],
    ["考试中心", "浏览 /exams/select", "题库选择与考试入口正常"],
    ["论坛与后台", "浏览 /forum、/admin/dashboard", "社区与后台统计页面可访问"],
]

TEST_SUMMARY = [
    "作品运行环境采用 Docker 开发模式，服务由 compose.dev.yml 管理。验证时重点关注容器服务状态、健康检查接口与关键业务页面是否可访问。",
    "在当前开发环境中，docker compose --env-file .env -f compose.dev.yml ps 显示 web、postgres、redis、worker 容器均处于 Up 状态；访问 /api/ping 与 /api/ping?deep=1 均返回 status=success，深度检查返回 db=true、redis=true，说明基础运行链路正常。",
    "从技术指标看，系统以 Docker + PostgreSQL + Redis 为基础设施，具备较好的部署一致性与扩展性；Web 端同时兼顾响应式布局、Session/JWT 双认证、CSRF/XHR 校验链和统一错误处理，满足演示、学习使用与后续迭代需要。",
]

INSTALL = [
    "准备 .env 文件，并根据环境填写必要的配置项，例如 WECHAT_APPID、WECHAT_SECRET、DASHSCOPE_API_KEY 等。",
    "在项目根目录执行：docker compose --env-file .env -f compose.dev.yml up。",
    "等待 web、postgres、redis、worker 容器启动完成后，执行 flask db upgrade 初始化数据库。",
    f"浏览器访问 {LOCAL_URL} 进行本地演示；正式展示可使用公网地址 {PUBLIC_URL}。",
    "推荐演示顺序为：首页 -> 题库广场 -> 题库名片详情 -> 我的题库 -> 数据中心 -> 考试中心 -> 论坛 -> 后台。",
    "如需联调小程序，可使用微信开发者工具打开 miniprogram-1/，并在 miniprogram-1/miniprogram/utils/config.ts 中切换 API 模式。",
]

SUMMARY = [
    "Sak-AI答题助手围绕“题库资源管理 + 学习闭环反馈 + 社区互动 + 运营支撑”构建完整 Web 应用，不仅实现了刷题功能，还将题库组织、考试训练、数据可视化、论坛交流和后台管理整合到同一系统中。",
    "工程上，作品采用单仓全栈架构，保证 Web 与小程序共享语义，降低维护成本；部署上，Docker 化运行方式提升了演示和复现效率；应用上，AI 题目解析能力让系统更适合面向教学与学习场景的持续扩展。后续可进一步增强推荐策略、协作编辑与个性化复盘能力。",
]

REFERENCES = [
    "Flask Official Documentation. https://flask.palletsprojects.com/",
    "SQLAlchemy Documentation. https://docs.sqlalchemy.org/",
    "Docker Docs. https://docs.docker.com/",
    "PostgreSQL Documentation. https://www.postgresql.org/docs/",
    "微信开放文档 - 小程序. https://developers.weixin.qq.com/miniprogram/dev/framework/",
]


def ensure_output_dir() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_run_font(run, *, font_name: str = "宋体", size: float = 10.5, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def replace_paragraph_text(paragraph: Paragraph, text: str, *, font_name: str = "宋体", size: float = 10.5, bold: bool = False) -> Paragraph:
    if paragraph.runs:
        paragraph.runs[0].text = text
        set_run_font(paragraph.runs[0], font_name=font_name, size=size, bold=bold)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(text)
        set_run_font(run, font_name=font_name, size=size, bold=bold)
    return paragraph


def copy_paragraph_format(src: Paragraph, dest: Paragraph) -> None:
    src_fmt = src.paragraph_format
    dest_fmt = dest.paragraph_format
    dest_fmt.left_indent = src_fmt.left_indent
    dest_fmt.right_indent = src_fmt.right_indent
    dest_fmt.first_line_indent = src_fmt.first_line_indent
    dest_fmt.keep_together = src_fmt.keep_together
    dest_fmt.keep_with_next = src_fmt.keep_with_next
    dest_fmt.page_break_before = src_fmt.page_break_before
    dest_fmt.widow_control = src_fmt.widow_control
    dest_fmt.space_before = src_fmt.space_before
    dest_fmt.space_after = src_fmt.space_after
    dest_fmt.line_spacing = src_fmt.line_spacing
    dest_fmt.line_spacing_rule = src_fmt.line_spacing_rule
    dest_fmt.alignment = src.alignment


def insert_paragraph_after(anchor: Paragraph | Table, *, text: str = "", font_name: str = "宋体", size: float = 10.5, bold: bool = False, align: WD_ALIGN_PARAGRAPH | None = None, copy_format_from: Paragraph | None = None) -> Paragraph:
    anchor_element = anchor._p if isinstance(anchor, Paragraph) else anchor._tbl
    parent = anchor._parent
    new_p = OxmlElement("w:p")
    anchor_element.addnext(new_p)
    paragraph = Paragraph(new_p, parent)
    paragraph.style = copy_format_from.style if copy_format_from is not None else "Normal"
    if copy_format_from is not None:
        copy_paragraph_format(copy_format_from, paragraph)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, font_name=font_name, size=size, bold=bold)
    return paragraph


def insert_picture_after(anchor: Paragraph | Table, image_path: Path, *, width_cm: float = 15.4, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER, copy_format_from: Paragraph | None = None) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, copy_format_from=copy_format_from)
    paragraph.alignment = align
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return paragraph


def fill_cell(cell, text: str, *, font_name: str = "宋体", size: float = 10.0, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, font_name=font_name, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def insert_table_after(document: Document, anchor: Paragraph | Table, rows: list[list[str]], widths_cm: list[float], *, font_size: float = 10.0) -> Table:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
    for row_index, row_values in enumerate(rows):
        for column_index, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            fill_cell(
                table.cell(row_index, column_index),
                value,
                size=font_size,
                bold=row_index == 0,
                align=align,
            )
    anchor_element = anchor._p if isinstance(anchor, Paragraph) else anchor._tbl
    anchor_element.addnext(table._tbl)
    return table


def insert_bullets(anchor: Paragraph | Table, bullets: list[str], *, copy_format_from: Paragraph) -> Paragraph | Table:
    current: Paragraph | Table = anchor
    for item in bullets:
        current = insert_paragraph_after(
            current,
            text=f"• {item}",
            copy_format_from=copy_format_from,
        )
    return current


def insert_numbered(anchor: Paragraph | Table, items: list[str], *, copy_format_from: Paragraph) -> Paragraph | Table:
    current: Paragraph | Table = anchor
    for index, item in enumerate(items, start=1):
        current = insert_paragraph_after(
            current,
            text=f"{index}. {item}",
            copy_format_from=copy_format_from,
        )
    return current


def fill_title_page(paragraphs: list[Paragraph]) -> None:
    replace_paragraph_text(paragraphs[5], f"作品编号：{PROJECT_NUMBER}", font_name="宋体", size=16)
    replace_paragraph_text(paragraphs[6], f"作品名称：{PROJECT_NAME}", font_name="宋体", size=16)
    replace_paragraph_text(paragraphs[7], f"作　　者：{AUTHORS}", font_name="宋体", size=16)
    replace_paragraph_text(paragraphs[8], f"版本编号：{VERSION}", font_name="宋体", size=16)
    replace_paragraph_text(paragraphs[9], f"填写日期：{DATE_TEXT}", font_name="宋体", size=16)


def fill_need_section(document: Document, base_paragraph: Paragraph) -> None:
    replace_paragraph_text(base_paragraph, NEEDS[0])
    current: Paragraph | Table = base_paragraph
    current = insert_paragraph_after(current, text=NEEDS[1], copy_format_from=base_paragraph)
    insert_table_after(document, current, COMPETITOR_ROWS, widths_cm=[3.2, 6.1, 6.1], font_size=9.5)


def fill_overview_section(document: Document, base_paragraph: Paragraph) -> None:
    replace_paragraph_text(base_paragraph, OVERVIEW[0])
    current: Paragraph | Table = base_paragraph
    for text in OVERVIEW[1:]:
        current = insert_paragraph_after(current, text=text, copy_format_from=base_paragraph)
    current = insert_bullets(current, OVERVIEW_BULLETS, copy_format_from=base_paragraph)
    current = insert_paragraph_after(
        current,
        text="图 1  系统整体架构图",
        font_name="宋体",
        size=10.5,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        copy_format_from=base_paragraph,
    )
    insert_picture_after(current, ARCH_IMG, copy_format_from=base_paragraph)


def fill_detail_section(document: Document, base_paragraph: Paragraph) -> None:
    replace_paragraph_text(base_paragraph, DETAIL[0])
    current: Paragraph | Table = base_paragraph
    current = insert_paragraph_after(
        current,
        text="图 2  核心用户流程图",
        font_name="宋体",
        size=10.5,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        copy_format_from=base_paragraph,
    )
    current = insert_picture_after(current, FLOW_IMG, copy_format_from=base_paragraph)
    current = insert_paragraph_after(current, text=DETAIL[1], copy_format_from=base_paragraph)
    current = insert_table_after(document, current, DATABASE_ROWS, widths_cm=[4.0, 11.4], font_size=9.5)
    current = insert_paragraph_after(current, text=DETAIL[2], copy_format_from=base_paragraph)
    current = insert_numbered(current, DETAIL_NUMBERED, copy_format_from=base_paragraph)
    current = insert_paragraph_after(
        current,
        text="真实运行界面截图",
        font_name="宋体",
        size=10.5,
        bold=True,
        copy_format_from=base_paragraph,
    )
    for index, (filename, route, description) in enumerate(SCREENSHOTS, start=3):
        current = insert_paragraph_after(
            current,
            text=f"图 {index}  {filename.removesuffix('.png')}（路径：{route}）",
            font_name="宋体",
            size=10.5,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            copy_format_from=base_paragraph,
        )
        current = insert_picture_after(current, SCREENSHOT_DIR / filename, width_cm=15.2, copy_format_from=base_paragraph)
        current = insert_paragraph_after(
            current,
            text=description,
            font_name="宋体",
            size=9.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            copy_format_from=base_paragraph,
        )


def fill_test_section(document: Document, base_paragraph: Paragraph) -> None:
    replace_paragraph_text(base_paragraph, TEST_SUMMARY[0])
    current: Paragraph | Table = base_paragraph
    current = insert_table_after(document, current, TEST_ROWS, widths_cm=[3.1, 6.0, 5.6], font_size=9.0)
    for text in TEST_SUMMARY[1:]:
        current = insert_paragraph_after(current, text=text, copy_format_from=base_paragraph)


def fill_install_section(base_paragraph: Paragraph) -> None:
    replace_paragraph_text(base_paragraph, f"1. {INSTALL[0]}")
    current: Paragraph | Table = base_paragraph
    for index, text in enumerate(INSTALL[1:], start=2):
        current = insert_paragraph_after(current, text=f"{index}. {text}", copy_format_from=base_paragraph)


def fill_summary_section(base_paragraph: Paragraph) -> None:
    replace_paragraph_text(base_paragraph, SUMMARY[0])
    insert_paragraph_after(base_paragraph, text=SUMMARY[1], copy_format_from=base_paragraph)


def fill_reference_section(base_paragraph: Paragraph) -> None:
    replace_paragraph_text(base_paragraph, f"1. {REFERENCES[0]}", font_name="宋体", size=10.5)
    current: Paragraph | Table = base_paragraph
    for index, text in enumerate(REFERENCES[1:], start=2):
        current = insert_paragraph_after(current, text=f"{index}. {text}", copy_format_from=base_paragraph)


def export_pdf_via_pages(docx_path: Path, pdf_path: Path) -> None:
    script = (
        'on run argv\n'
        '  set inputPath to POSIX file (item 1 of argv)\n'
        '  set outputPath to POSIX file (item 2 of argv)\n'
        '  tell application "Pages"\n'
        '    activate\n'
        '    set theDoc to open inputPath\n'
        '    export theDoc to outputPath as PDF\n'
        '    close theDoc saving no\n'
        '    quit\n'
        '  end tell\n'
        'end run\n'
    )
    script_path = OUTPUT_DIR / "export_design_doc_to_pdf.applescript"
    script_path.write_text(script, encoding="utf-8")
    import subprocess

    subprocess.run(
        ["osascript", str(script_path), str(docx_path), str(pdf_path)],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )


def main() -> None:
    ensure_output_dir()
    document = Document(str(TEMPLATE_PATH))
    original_paragraphs = list(document.paragraphs)
    fill_title_page(original_paragraphs)
    fill_need_section(document, original_paragraphs[16])
    fill_overview_section(document, original_paragraphs[18])
    fill_detail_section(document, original_paragraphs[20])
    fill_test_section(document, original_paragraphs[22])
    fill_install_section(original_paragraphs[24])
    fill_summary_section(original_paragraphs[26])
    fill_reference_section(original_paragraphs[28])
    document.save(str(OUTPUT_DOCX))
    document.save(str(DOWNLOAD_DOCX))
    export_pdf_via_pages(OUTPUT_DOCX, OUTPUT_PDF)


if __name__ == "__main__":
    main()
