# -*- coding: utf-8 -*-
"""个人题库 Word/PDF 导出测试。"""

from __future__ import annotations

from io import BytesIO
import json
import uuid
import zipfile

from docx import Document
from sqlalchemy import text

from app.core.extensions import db


def _create_bank_with_questions(app, user_id: int) -> tuple[int, dict[str, int]]:
    bank_name = f"文档导出测试题库-{uuid.uuid4().hex[:8]}"
    with app.app_context():
        bank_id = db.session.execute(
            text(
                """
                INSERT INTO user_question_banks (user_id, name, status, question_count)
                VALUES (:user_id, :name, 1, 2)
                RETURNING id
                """
            ),
            {"user_id": int(user_id), "name": bank_name},
        ).scalar_one()

        q1_id = db.session.execute(
            text(
                """
                INSERT INTO user_bank_questions (
                    bank_id, user_id, type, content, options, answer,
                    analysis, tags, difficulty, source_type, sort_order
                )
                VALUES (
                    :bank_id, :user_id, 'single_choice', :content, :options,
                    :answer, :analysis, '[]', 1, 'custom', 1
                )
                RETURNING id
                """
            ),
            {
                "bank_id": int(bank_id),
                "user_id": int(user_id),
                "content": "导出选择题题干",
                "options": json.dumps(["甲选项", "乙选项"], ensure_ascii=False),
                "answer": json.dumps([0], ensure_ascii=False),
                "analysis": "导出选择题解析",
            },
        ).scalar_one()

        q2_id = db.session.execute(
            text(
                """
                INSERT INTO user_bank_questions (
                    bank_id, user_id, type, content, options, answer,
                    analysis, tags, difficulty, source_type, sort_order
                )
                VALUES (
                    :bank_id, :user_id, 'fill', :content, '[]',
                    :answer, :analysis, '[]', 2, 'custom', 2
                )
                RETURNING id
                """
            ),
            {
                "bank_id": int(bank_id),
                "user_id": int(user_id),
                "content": "导出填空题题干：{0}",
                "answer": json.dumps([["标准答案"]], ensure_ascii=False),
                "analysis": "导出填空题解析",
            },
        ).scalar_one()

        db.session.commit()
        return int(bank_id), {"choice": int(q1_id), "fill": int(q2_id)}


def _delete_bank(app, bank_id: int | None) -> None:
    if bank_id is None:
        return
    with app.app_context():
        db.session.execute(text("DELETE FROM user_bank_questions WHERE bank_id = :bank_id"), {"bank_id": int(bank_id)})
        db.session.execute(text("DELETE FROM user_question_banks WHERE id = :bank_id"), {"bank_id": int(bank_id)})
        db.session.commit()


def _docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    chunks: list[str] = []
    chunks.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    return "\n".join(chunks)


def test_manage_page_offers_word_and_pdf_export_actions(app, auth_client, seed_user):
    bank_id, _ids = _create_bank_with_questions(app, seed_user["id"])

    try:
        response = auth_client.get(f"/user/banks/{bank_id}")

        assert response.status_code == 200
        html = response.get_data(as_text=True)
        assert "导出 Word" in html
        assert "导出 PDF" in html
        assert f"/user/banks/api/{bank_id}/questions/export/word" in html
        assert f"/user/banks/api/{bank_id}/questions/export/pdf" in html
        assert 'id="ubmBatchExportWordBtn"' in html
        assert 'id="ubmBatchExportPdfBtn"' in html
    finally:
        _delete_bank(app, bank_id)


def test_export_word_includes_promo_qr_and_question_content(app, auth_client, seed_user):
    bank_id, _ids = _create_bank_with_questions(app, seed_user["id"])

    try:
        response = auth_client.get(f"/user/banks/api/{bank_id}/questions/export/word")

        assert response.status_code == 200
        assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert ".docx" in response.headers.get("Content-Disposition", "")

        data = response.get_data()
        text_content = _docx_text(data)
        assert "https://www.saksk.top" in text_content
        assert "扫码打开小程序" in text_content
        assert "导出选择题题干" in text_content
        assert "甲选项" in text_content
        assert "答案：A" in text_content
        assert "导出填空题题干" in text_content

        with zipfile.ZipFile(BytesIO(data)) as zf:
            assert any(name.startswith("word/media/") for name in zf.namelist())
    finally:
        _delete_bank(app, bank_id)


def test_selected_word_export_only_includes_selected_questions(app, auth_client, seed_user):
    bank_id, ids = _create_bank_with_questions(app, seed_user["id"])

    try:
        response = auth_client.get(
            f"/user/banks/api/{bank_id}/questions/export/word?ids={ids['fill']}"
        )

        assert response.status_code == 200
        text_content = _docx_text(response.get_data())
        assert "导出填空题题干" in text_content
        assert "标准答案" in text_content
        assert "导出选择题题干" not in text_content
    finally:
        _delete_bank(app, bank_id)


def test_export_pdf_returns_pdf_download(app, auth_client, seed_user):
    bank_id, _ids = _create_bank_with_questions(app, seed_user["id"])

    try:
        response = auth_client.get(f"/user/banks/api/{bank_id}/questions/export/pdf")

        assert response.status_code == 200
        assert response.mimetype == "application/pdf"
        assert ".pdf" in response.headers.get("Content-Disposition", "")
        assert response.get_data().startswith(b"%PDF")
    finally:
        _delete_bank(app, bank_id)
