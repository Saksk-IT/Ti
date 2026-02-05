# -*- coding: utf-8 -*-
"""Admin API routes - questions import/export."""

import datetime
import io
import json
import os
import sqlite3
import zipfile

import pandas as pd
from flask import (
    Blueprint,
    current_app,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    send_from_directory,
    session,
    url_for,
)
from werkzeug.security import generate_password_hash
from werkzeug.utils import secure_filename

from app.core.extensions import limiter
from app.core.utils.database import get_db
from app.core.utils.cache_utils import bump_questions_version, bump_subjects_version
from app.core.utils.fill_blank_parser import parse_fill_blank
from app.core.utils.validators import parse_int, validate_password

from ..api_bp import admin_api_bp


@admin_api_bp.route('/questions/import', methods=['POST'])
def import_questions_api():
    """导入题目"""
    data = request.json
    subject_id = data.get('subject_id')
    questions = data.get('questions', [])
    
    if not subject_id or not questions:
        return jsonify({'status': 'error', 'message': '缺少科目或题库数据'}), 400
    
    conn = get_db()
    try:
        q_cols = [r['name'] for r in conn.execute("PRAGMA table_info(questions)").fetchall()]
    except Exception:
        q_cols = []
    has_difficulty = 'difficulty' in q_cols
    has_tags = 'tags' in q_cols
    has_created_by = 'created_by' in q_cols
    has_updated_at = 'updated_at' in q_cols

    from app.core.utils.options_parser import parse_options
    from app.core.utils.portable_question_format import portable_question_to_internal, tags_to_storage_str, internal_question_to_portable

    imported = 0
    try:
        for item in questions:
            if not isinstance(item, dict):
                continue

            # 新统一格式：{type, content, options, answer, analysis, tags, difficulty}
            if 'type' in item or (isinstance(item.get('answer'), list) and 'content' in item):
                internal, conv_errors = portable_question_to_internal(item, scope='question_center')
                if conv_errors:
                    continue

                q_type = internal.get('q_type') or '未知'
                content = internal.get('content') or ''
                answer = internal.get('answer') or ''
                explanation = internal.get('explanation') or ''
                options_list = internal.get('options') or []
                opts_json = json.dumps(options_list, ensure_ascii=False) if options_list else '[]'
                diff_val = int(internal.get('difficulty') or 1)
                tags_str = tags_to_storage_str(internal.get('tags') or [])
            else:
                # 兼容旧格式（题型/题干/选项/答案/解析/难度/标签）
                q_type = item.get('题型', '未知')
                content = item.get('题干', '')
                answer = item.get('答案', '')
                explanation = item.get('解析', '')
                diff_val = item.get('难度', 1)
                tags_str = tags_to_storage_str(item.get('标签') or item.get('tags') or [])

                options_list = item.get('选项') or item.get('options') or []
                if isinstance(options_list, str):
                    try:
                        options_list = json.loads(options_list)
                    except Exception:
                        options_list = []
                if isinstance(options_list, list):
                    options_list = [str(o) for o in options_list]
                else:
                    options_list = []
                opts_json = json.dumps(options_list, ensure_ascii=False) if options_list else '[]'

            if q_type in ('选择题', '多选题'):
                if q_type == '多选题':
                    answer = (answer or '').strip()
                    if len(answer) < 2:
                        continue
                    try:
                        parsed_options = parse_options(options_list)
                        valid_keys = {opt['key'] for opt in parsed_options if opt.get('key')}
                        invalid_keys = set(answer.upper()) - valid_keys
                        if invalid_keys:
                            continue
                    except Exception:
                        pass
            elif q_type == '填空题':
                new_content, new_answer, _blank_count = parse_fill_blank(content)
                if new_answer:
                    content = new_content
                    answer = new_answer

            # 统一落库为 PQF 同名列
            portable = internal_question_to_portable(
                q_id=None,
                q_type=q_type,
                content=content,
                options=opts_json,
                answer=answer,
                explanation=explanation,
                difficulty=diff_val,
                tags=tags_str,
            )

            cols = ['subject_id', 'type', 'content', 'options', 'answer', 'analysis']
            vals = [
                subject_id,
                portable.get('type') or 'essay',
                portable.get('content') or '',
                json.dumps(portable.get('options') or [], ensure_ascii=False),
                json.dumps(portable.get('answer') if portable.get('answer') is not None else [], ensure_ascii=False),
                portable.get('analysis') or '',
            ]

            if has_difficulty:
                cols.append('difficulty')
                vals.append(int(portable.get('difficulty') or 1))
            if has_tags:
                cols.append('tags')
                vals.append(json.dumps(portable.get('tags') or [], ensure_ascii=False))
            if has_created_by:
                cols.append('created_by')
                vals.append(session.get('user_id'))

            # 可选：图片字段（旧导入可能带了）
            if 'image_path' in q_cols:
                cols.append('image_path')
                vals.append(item.get('image_path'))

            if has_updated_at:
                sql = f"INSERT INTO questions ({', '.join(cols)}, updated_at) VALUES ({', '.join(['?'] * len(vals))}, CURRENT_TIMESTAMP)"
            else:
                sql = f"INSERT INTO questions ({', '.join(cols)}) VALUES ({', '.join(['?'] * len(vals))})"

            conn.execute(sql, vals)
            imported += 1

        conn.commit()
        try:
            bump_questions_version()
        except Exception:
            pass
        return jsonify({'status': 'success', 'message': f'成功导入{imported}道题'})
    except Exception as e:
        try:
            conn.rollback()
        except Exception:
            pass
        return jsonify({'status': 'error', 'message': str(e)})



@admin_api_bp.route('/questions/export', methods=['GET'])
def export_questions_api():
    """导出题目"""
    subject_id = request.args.get('subject_id')
    
    conn = get_db()
    try:
        q_cols = [r['name'] for r in conn.execute("PRAGMA table_info(questions)").fetchall()]
    except Exception:
        q_cols = []
    has_difficulty = 'difficulty' in q_cols
    has_tags = 'tags' in q_cols

    select_cols = [
        'q.id',
        'q.subject_id',
        's.name as subject_name',
        'q.content',
        'q.type',
        'q.options',
        'q.answer',
        'q.analysis',
    ]
    if has_difficulty:
        select_cols.append('q.difficulty')
    if has_tags:
        select_cols.append('q.tags')

    sql = '''
        SELECT {cols}
        FROM questions q
        LEFT JOIN subjects s ON q.subject_id = s.id
        WHERE 1=1
    '''.format(cols=', '.join(select_cols))
    params = []
    
    if subject_id:
        sql += ' AND q.subject_id = ?'
        params.append(subject_id)
    
    sql += ' ORDER BY q.id'
    rows = conn.execute(sql, params).fetchall()
    
    def _safe_load(raw, default):
        if raw is None:
            return default
        if isinstance(raw, (list, dict, bool, int, float)):
            return raw
        s = str(raw).strip()
        if not s:
            return default
        try:
            return json.loads(s)
        except Exception:
            return default

    items = []
    for r in rows:
        item = {
            'id': int(r['id']),
            'type': (r['type'] or ''),
            'content': (r['content'] or ''),
            'options': _safe_load(r['options'], []),
            'answer': _safe_load(r['answer'], []),
            'analysis': (r['analysis'] or ''),
            'tags': _safe_load(r['tags'], []) if has_tags else [],
            'difficulty': int(r['difficulty'] or 1) if has_difficulty else 1,
        }
        # 附带科目信息（导出全量备份时有用；导入可忽略）
        item['subject_id'] = r['subject_id']
        item['subject_name'] = r['subject_name'] or '默认科目'
        items.append(item)

    meta = {'scope': 'question_center'}
    if subject_id:
        meta['subject_id'] = subject_id
    return jsonify({'status': 'success', 'meta': meta, 'count': len(items), 'questions': items})

@admin_api_bp.route('/questions/import/excel', methods=['POST'])
def import_questions_from_excel():
    """从Excel文件导入题库 (V2 - 分列格式)"""
    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': '没有文件部分'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'status': 'error', 'message': '没有选择文件'}), 400

    if not file or not file.filename.endswith('.xlsx'):
        return jsonify({'status': 'error', 'message': '只允许上传 .xlsx 文件'}), 400

    conn = get_db()
    
    try:
        from app.core.utils.portable_question_format import internal_question_to_portable

        df = pd.read_excel(file, sheet_name='题目示例').fillna('')
        
        required_columns = ['subject', 'q_type', 'content']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            return jsonify({'status': 'error', 'message': f'Excel文件中缺少必需的列: {", ".join(missing_columns)}'}), 400

        subjects = conn.execute('SELECT id, name FROM subjects').fetchall()
        subject_map = {s['name']: s['id'] for s in subjects}

        imported_count = 0
        errors = []

        option_cols = sorted([col for col in df.columns if col.startswith('option_')])
        blank_cols = sorted([col for col in df.columns if col.startswith('blank_')])

        for index, row in df.iterrows():
            try:
                subject_name = str(row.get('subject', '')).strip()
                q_type = str(row.get('q_type', '')).strip()
                content = str(row.get('content', '')).strip()
                answer = str(row.get('answer', '')).strip()
                explanation = str(row.get('explanation', '')).strip()

                if not all([subject_name, q_type, content]):
                    errors.append(f'第 {index + 2} 行: 必填字段（subject, q_type, content）不能为空。')
                    continue

                if subject_name not in subject_map:
                    cursor = conn.execute('INSERT INTO subjects (name) VALUES (?)', (subject_name,))
                    subject_id = cursor.lastrowid
                    subject_map[subject_name] = subject_id
                    conn.commit()
                    try:
                        bump_subjects_version()
                    except Exception:
                        pass
                else:
                    subject_id = subject_map[subject_name]

                valid_q_types = ['选择题', '多选题', '判断题', '填空题', '简答题']
                if q_type not in valid_q_types:
                    errors.append(f'第 {index + 2} 行: 无效的题型 "{q_type}"。')
                    continue
                
                options_json = '[]'
                final_answer = answer

                if q_type in ['选择题', '多选题']:
                    options_list = []
                    for i, col_name in enumerate(option_cols):
                        option_text = str(row.get(col_name, '')).strip()
                        if option_text:
                            prefix = chr(ord('A') + i)
                            options_list.append(f"{prefix}. {option_text}")
                    if not options_list:
                        errors.append(f'第 {index + 2} 行: 选择题或多选题至少需要一个选项。')
                        continue
                    options_json = json.dumps(options_list, ensure_ascii=False)
                    if not final_answer:
                        errors.append(f'第 {index + 2} 行: 选择题或多选题的 `answer` 列不能为空。')
                        continue

                elif q_type == '填空题':
                    blank_answers = []
                    for col_name in blank_cols:
                        blank_text = str(row.get(col_name, '')).strip()
                        if blank_text:
                            blank_answers.append(blank_text)
                    if not blank_answers:
                        errors.append(f'第 {index + 2} 行: 填空题至少需要一个 `blank_` 答案。')
                        continue
                    final_answer = ';;'.join(blank_answers)
                
                elif not final_answer:
                     errors.append(f'第 {index + 2} 行: `answer` 列不能为空。')
                     continue

                portable = internal_question_to_portable(
                    q_id=None,
                    q_type=q_type,
                    content=content,
                    options=options_json,
                    answer=final_answer,
                    explanation=explanation,
                    difficulty=1,
                    tags='',
                )

                conn.execute(
                    '''
                    INSERT INTO questions
                    (subject_id, type, content, options, answer, analysis, difficulty, tags, created_by, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
                    ''',
                    (
                        subject_id,
                        portable.get('type') or 'essay',
                        portable.get('content') or '',
                        json.dumps(portable.get('options') or [], ensure_ascii=False),
                        json.dumps(portable.get('answer') if portable.get('answer') is not None else [], ensure_ascii=False),
                        portable.get('analysis') or '',
                        int(portable.get('difficulty') or 1),
                        json.dumps(portable.get('tags') or [], ensure_ascii=False),
                        session.get('user_id'),
                    ),
                )
                
                imported_count += 1

            except Exception as e:
                errors.append(f'第 {index + 2} 行: 导入失败 - {str(e)}')

        conn.commit()
        try:
            bump_questions_version()
        except Exception:
            pass
        try:
            bump_subjects_version()
        except Exception:
            pass
        
        message = f'成功导入 {imported_count} 道题。'
        if errors:
            message += f' 遇到 {len(errors)} 个问题。'
        
        return jsonify({
            'status': 'success' if not errors else 'warning',
            'message': message,
            'imported_count': imported_count,
            'errors': errors
        })

    except Exception as e:
        return jsonify({'status': 'error', 'message': f'处理文件失败: {str(e)}'}), 500



@admin_api_bp.route('/download_template')
def download_template():
    """提供题库导入模板文件的下载"""
    directory = os.path.join(current_app.root_path, '..', 'instance')
    return send_from_directory(directory, 'question_import_template.xlsx', as_attachment=True)



@admin_api_bp.route('/questions/export/excel', methods=['GET'])
def export_questions_to_excel():
    """导出题目为Excel文件（使用与导入相同的模板格式）"""
    subject_id = request.args.get('subject_id')
    q_type = request.args.get('type', 'all')
    
    conn = get_db()
    
    # 构建查询SQL
    sql = '''
        SELECT q.*, s.name as subject_name
        FROM questions q
        LEFT JOIN subjects s ON q.subject_id = s.id
        WHERE 1=1
    '''
    params = []
    
    if subject_id:
        sql += ' AND q.subject_id = ?'
        params.append(subject_id)
    
    if q_type and q_type != 'all':
        from app.core.utils.portable_question_format import any_type_to_portable_type

        sql += ' AND q.type = ?'
        params.append(any_type_to_portable_type(q_type))
    
    sql += ' ORDER BY q.id'
    rows = conn.execute(sql, params).fetchall()
    
    if not rows:
        return jsonify({'status': 'error', 'message': '没有可导出的题目'}), 400

    from app.core.utils.portable_question_format import portable_question_to_internal

    def _safe_load(raw, default):
        if raw is None:
            return default
        if isinstance(raw, (list, dict, bool, int, float)):
            return raw
        s = str(raw).strip()
        if not s:
            return default
        try:
            return json.loads(s)
        except Exception:
            return default

    # 准备数据
    export_data = []
    max_options = 0
    max_blanks = 0
    seed = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

    for row in rows:
        question = dict(row)
        portable = {
            'id': question.get('id'),
            'type': question.get('type') or '',
            'content': question.get('content') or '',
            'options': _safe_load(question.get('options'), []),
            'answer': _safe_load(question.get('answer'), []),
            'analysis': question.get('analysis') or '',
            'tags': _safe_load(question.get('tags'), []),
            'difficulty': int(question.get('difficulty') or 1),
        }
        internal, _errors = portable_question_to_internal(portable, scope='question_center')

        q_type_val = internal.get('q_type') or ''
        content = internal.get('content') or ''
        answer_text = internal.get('answer') or ''
        explanation = internal.get('explanation') or ''
        options = internal.get('options') or []

        blank_answers = []
        if q_type_val == '填空题' and answer_text:
            blank_answers = str(answer_text).split(';;')

        max_options = max(max_options, len(options))
        max_blanks = max(max_blanks, len(blank_answers))

        row_data = {
            'subject': question.get('subject_name') or '默认科目',
            'q_type': q_type_val,
            'content': content,
            'answer': answer_text if q_type_val != '填空题' else '',
            'explanation': explanation,
        }

        # 添加选项列（移除 A./B. 前缀，保持与导入模板一致）
        for i, opt in enumerate(options):
            opt_text = str(opt or '')
            prefix = chr(ord('A') + i) + '. '
            if opt_text.startswith(prefix):
                opt_text = opt_text[len(prefix):]
            col = f"option_{seed[i]}" if i < len(seed) else f"option_{i+1}"
            row_data[col] = opt_text

        for i, blank in enumerate(blank_answers):
            row_data[f'blank_{i+1}'] = blank

        export_data.append(row_data)

    # 构建DataFrame
    columns = ['subject', 'q_type', 'content']

    # 添加选项列
    for i in range(max_options):
        columns.append(f"option_{seed[i]}" if i < len(seed) else f"option_{i+1}")

    columns.append('answer')

    # 添加填空题答案列
    for i in range(max_blanks):
        columns.append(f'blank_{i+1}')

    columns.append('explanation')

    # 创建DataFrame
    df = pd.DataFrame(export_data)
    
    # 确保所有列都存在
    for col in columns:
        if col not in df.columns:
            df[col] = ''
    
    # 按列顺序重新排列
    df = df[columns]
    
    # 创建Excel文件
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='题目示例', index=False)
    
    output.seek(0)
    
    # 生成文件名
    subject_name = "all_subjects"
    if subject_id:
        subject_row = conn.execute('SELECT name FROM subjects WHERE id = ?', (subject_id,)).fetchone()
        if subject_row:
            subject_name = subject_row['name']
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"questions_export_{subject_name}_{timestamp}.xlsx"
    
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )



@admin_api_bp.route('/questions/export/word', methods=['GET'])
def export_questions_to_word():
    """导出题目为Word文档"""
    # 延迟导入，避免模块加载时出错
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        return jsonify({'status': 'error', 'message': f'Word导出功能需要安装python-docx库: {str(e)}'}), 500
    
    subject_id = request.args.get('subject_id')
    q_type = request.args.get('type', 'all')
    
    conn = get_db()
    
    # 构建查询SQL
    sql = '''
        SELECT q.*, s.name as subject_name
        FROM questions q
        LEFT JOIN subjects s ON q.subject_id = s.id
        WHERE 1=1
    '''
    params = []
    
    if subject_id:
        sql += ' AND q.subject_id = ?'
        params.append(subject_id)
    
    if q_type and q_type != 'all':
        from app.core.utils.portable_question_format import any_type_to_portable_type

        sql += ' AND q.type = ?'
        params.append(any_type_to_portable_type(q_type))
    
    sql += ' ORDER BY q.id'
    rows = conn.execute(sql, params).fetchall()
    
    if not rows:
        return jsonify({'status': 'error', 'message': '没有可导出的题目'}), 400
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档标题样式
    title = doc.add_heading('题目导出', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加科目信息
    subject_name = "所有科目"
    if subject_id:
        subject_row = conn.execute('SELECT name FROM subjects WHERE id = ?', (subject_id,)).fetchone()
        if subject_row:
            subject_name = subject_row['name']
    
    info_para = doc.add_paragraph(f'科目：{subject_name}')
    if q_type and q_type != 'all':
        info_para.add_run(f' | 题型：{q_type}')
    info_para.add_run(f' | 题目数量：{len(rows)}')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分隔线
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    from app.core.utils.portable_question_format import portable_question_to_internal

    def _safe_load(raw, default):
        if raw is None:
            return default
        if isinstance(raw, (list, dict, bool, int, float)):
            return raw
        s = str(raw).strip()
        if not s:
            return default
        try:
            return json.loads(s)
        except Exception:
            return default
    
    # 遍历题目，添加到文档
    for idx, row in enumerate(rows, 1):
        question = dict(row)
        portable = {
            'id': question.get('id'),
            'type': question.get('type') or '',
            'content': question.get('content') or '',
            'options': _safe_load(question.get('options'), []),
            'answer': _safe_load(question.get('answer'), []),
            'analysis': question.get('analysis') or '',
            'tags': _safe_load(question.get('tags'), []),
            'difficulty': int(question.get('difficulty') or 1),
        }
        internal, _errors = portable_question_to_internal(portable, scope='question_center')

        q_type_val = internal.get('q_type') or ''
        content = internal.get('content') or ''
        answer = internal.get('answer') or ''
        explanation = internal.get('explanation') or ''
        
        # 添加题号
        q_num = doc.add_heading(f'题目 {idx}', level=2)
        q_num_format = q_num.runs[0].font
        q_num_format.size = Pt(14)
        q_num_format.bold = True
        
        # 添加题型标签
        type_para = doc.add_paragraph()
        type_run = type_para.add_run(f'【{q_type_val}】')
        type_run.font.bold = True
        type_run.font.color.rgb = RGBColor(0, 102, 204)
        
        # 添加题干
        content_para = doc.add_paragraph()
        content_run = content_para.add_run('题干：')
        content_run.font.bold = True
        content_para.add_run(content)
        
        # 解析并添加选项（如果是选择题或多选题）
        options = portable.get('options') or []
        
        if options:
            options_para = doc.add_paragraph()
            options_run = options_para.add_run('选项：')
            options_run.font.bold = True
            doc.add_paragraph()  # 空行
            
            for i, opt in enumerate(options):
                opt_para = doc.add_paragraph(f'{chr(ord("A") + i)}. {opt}', style='List Bullet')
                opt_para.paragraph_format.left_indent = Inches(0.5)
        
        # 添加答案
        answer_para = doc.add_paragraph()
        answer_run = answer_para.add_run('答案：')
        answer_run.font.bold = True
        answer_para.add_run(answer)
        
        # 如果是填空题，格式化显示答案
        if q_type_val == '填空题' and answer:
            blank_answers = answer.split(';;')
            if len(blank_answers) > 1:
                answer_para.clear()
                answer_run = answer_para.add_run('答案：')
                answer_run.font.bold = True
                for i, blank in enumerate(blank_answers, 1):
                    if i > 1:
                        answer_para.add_run(' | ')
                    answer_para.add_run(f'空{i}: {blank}')
        
        # 添加解析（如果有）
        if explanation:
            explanation_para = doc.add_paragraph()
            explanation_run = explanation_para.add_run('解析：')
            explanation_run.font.bold = True
            explanation_para.add_run(explanation)
        
        # 添加分隔线
        if idx < len(rows):
            doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # 空行
    
    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # 生成文件名
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"questions_export_{subject_name}_{timestamp}.docx"
    
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )



@admin_api_bp.route('/questions/upload_image', methods=['POST'])
def upload_question_image():
    """上传题目图片"""
    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': '没有文件部分'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'status': 'error', 'message': '没有选择文件'}), 400

    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'svg'}
    if '.' not in file.filename or file.filename.rsplit('.', 1)[1].lower() not in allowed_extensions:
        return jsonify({'status': 'error', 'message': '无效的文件类型'}), 400

    try:
        filename = secure_filename(file.filename)
        # 为了避免重名，可以加上时间戳和随机数
        import time, random
        unique_filename = f"{int(time.time())}_{random.randint(1000, 9999)}_{filename}"
        
        upload_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'question_images')
        file_path = os.path.join(upload_path, unique_filename)
        file.save(file_path)
        
        # 返回可访问的URL
        file_url = url_for('main.main_pages.serve_upload', filename=f'question_images/{unique_filename}')
        
        return jsonify({'status': 'success', 'url': file_url, 'path': f'question_images/{unique_filename}'})

    except Exception as e:
        return jsonify({'status': 'error', 'message': f'上传失败: {str(e)}'}), 500



@admin_api_bp.route('/questions/export_package', methods=['GET'])
def export_questions_package():
    """导出包含完整数据和图片的题目包"""
    subject_id = request.args.get('subject_id')
    q_type = request.args.get('type')
    
    conn = get_db()
    from app.core.utils.portable_question_format import any_type_to_portable_type
    
    # 1. 获取科目名称
    subject_name = "all_subjects"
    if subject_id:
        subject_row = conn.execute('SELECT name FROM subjects WHERE id = ?', (subject_id,)).fetchone()
        if subject_row:
            subject_name = subject_row['name']

    def _normalize_image_paths(raw_val):
        if raw_val is None:
            return []
        if isinstance(raw_val, list):
            return [str(x).strip() for x in raw_val if str(x).strip()]
        s = str(raw_val or '').strip()
        if not s or s in ('[]', '[ ]'):
            return []
        try:
            parsed = json.loads(s)
            if isinstance(parsed, list):
                return [str(x).strip() for x in parsed if str(x).strip()]
        except Exception:
            pass
        return [s]

    # 2. 查询题目数据
    sql = '''
        SELECT q.id, q.subject_id, s.name as subject_name,
               q.type, q.content, q.options, q.answer, q.analysis,
               q.difficulty, q.tags, q.image_path
        FROM questions q
        LEFT JOIN subjects s ON q.subject_id = s.id
        WHERE 1=1
    '''
    params = []
    if subject_id:
        sql += ' AND q.subject_id = ?'
        params.append(subject_id)
    if q_type and q_type != 'all':
        sql += ' AND q.type = ?'
        params.append(any_type_to_portable_type(q_type))

    sql += ' ORDER BY q.id'
    rows = conn.execute(sql, params).fetchall()

    # 3. 创建 ZIP 文件（统一 Portable Question Format）
    memory_file = io.BytesIO()
    upload_folder = current_app.config.get('UPLOAD_FOLDER', os.path.join(current_app.root_path, '..', 'uploads'))
    questions_data = []

    def _safe_load(raw, default):
        if raw is None:
            return default
        if isinstance(raw, (list, dict, bool, int, float)):
            return raw
        s = str(raw).strip()
        if not s:
            return default
        try:
            return json.loads(s)
        except Exception:
            return default

    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        for r in rows:
            item = {
                'id': int(r['id']),
                'type': (r['type'] or ''),
                'content': (r['content'] or ''),
                'options': _safe_load(r['options'], []),
                'answer': _safe_load(r['answer'], []),
                'analysis': (r['analysis'] or ''),
                'tags': _safe_load(r['tags'], []),
                'difficulty': int(r['difficulty'] or 1),
            }
            item['subject_id'] = r['subject_id']
            item['subject_name'] = r['subject_name'] or '默认科目'

            images_in_zip = []
            for img_rel in _normalize_image_paths(r['image_path']):
                img_rel = str(img_rel).replace('\\', '/').lstrip('/')
                full_image_path = os.path.join(upload_folder, *img_rel.split('/'))
                if not os.path.exists(full_image_path):
                    continue
                arcname = f"images/{img_rel}"
                try:
                    zf.write(full_image_path, arcname)
                    images_in_zip.append(arcname)
                except Exception:
                    continue
            item['images'] = images_in_zip
            questions_data.append(item)

        payload = {
            'meta': {
                'scope': 'question_center_package',
                'subject_id': int(subject_id) if str(subject_id or '').isdigit() else None,
                'subject_name': subject_name,
                'exported_at': datetime.datetime.now().isoformat(),
            },
            'questions': questions_data,
        }
        zf.writestr('data.json', json.dumps(payload, ensure_ascii=False, indent=2))

    memory_file.seek(0)
    
    # 4. 生成文件名并发送文件
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"questions_export_{subject_name}_{timestamp}.zip"
    
    return send_file(
        memory_file,
        as_attachment=True,
        download_name=filename,
        mimetype='application/zip'
    )



@admin_api_bp.route('/questions/import_package', methods=['POST'])
def import_questions_package():
    """导入题目包 (.zip)"""
    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': '没有文件部分'}), 400

    file = request.files['file']
    if file.filename == '' or not file.filename.endswith('.zip'):
        return jsonify({'status': 'error', 'message': '请上传有效的 .zip 文件'}), 400

    conn = get_db()
    
    # 获取现有的科目 name -> id 映射
    subjects = conn.execute('SELECT id, name FROM subjects').fetchall()
    subject_map = {s['name']: s['id'] for s in subjects}

    imported_count = 0
    errors = []
    upload_folder = current_app.config.get('UPLOAD_FOLDER', os.path.join(current_app.root_path, '..', 'uploads'))

    try:
        import uuid
        from app.core.utils.portable_question_format import (
            internal_question_to_portable,
            portable_question_to_internal,
            tags_to_storage_str,
        )

        try:
            q_cols = [r['name'] for r in conn.execute("PRAGMA table_info(questions)").fetchall()]
        except Exception:
            q_cols = []
        has_difficulty = 'difficulty' in q_cols
        has_tags = 'tags' in q_cols
        has_image_path = 'image_path' in q_cols
        has_created_by = 'created_by' in q_cols
        has_updated_at = 'updated_at' in q_cols

        def _normalize_image_paths(raw_val):
            if raw_val is None:
                return []
            if isinstance(raw_val, list):
                return [str(x).strip() for x in raw_val if str(x).strip()]
            s = str(raw_val or '').strip()
            if not s or s in ('[]', '[ ]'):
                return []
            try:
                parsed = json.loads(s)
                if isinstance(parsed, list):
                    return [str(x).strip() for x in parsed if str(x).strip()]
            except Exception:
                pass
            return [s]

        with zipfile.ZipFile(file, 'r') as zf:
            if 'data.json' not in zf.namelist():
                return jsonify({'status': 'error', 'message': '压缩包中缺少 data.json 文件'}), 400
            
            with zf.open('data.json') as f:
                raw_payload = json.load(f)

            meta = raw_payload.get('meta', {}) if isinstance(raw_payload, dict) else {}
            if isinstance(raw_payload, dict):
                questions_data = raw_payload.get('questions') or []
            elif isinstance(raw_payload, list):
                questions_data = raw_payload
            else:
                return jsonify({'status': 'error', 'message': 'data.json 格式不支持'}), 400

            default_subject_name = meta.get('subject_name') if isinstance(meta, dict) else None
            
            for q in questions_data:
                try:
                    if not isinstance(q, dict):
                        continue

                    # 1. 处理科目
                    subject_name = q.get('subject_name') or default_subject_name
                    if not subject_name:
                        errors.append(f"题目ID {q.get('id')} 缺少科目名称，已跳过。")
                        continue

                    if subject_name not in subject_map:
                        cursor = conn.execute('INSERT INTO subjects (name) VALUES (?)', (subject_name,))
                        subject_id = cursor.lastrowid
                        subject_map[subject_name] = subject_id
                    else:
                        subject_id = subject_map[subject_name]

                    # 2. 解析题目（统一格式优先）
                    if 'type' in q or (isinstance(q.get('answer'), list) and 'content' in q):
                        internal, conv_errors = portable_question_to_internal(q, scope='question_center')
                        if conv_errors:
                            errors.append(f"题目ID {q.get('id', 'N/A')} 导入失败: {'；'.join(conv_errors)}")
                            continue
                        q_type = internal.get('q_type') or '未知'
                        content = internal.get('content') or ''
                        answer = internal.get('answer') or ''
                        explanation = internal.get('explanation') or ''
                        options_json = json.dumps(internal.get('options') or [], ensure_ascii=False)
                        diff_val = int(internal.get('difficulty') or 1)
                        tags_str = tags_to_storage_str(internal.get('tags') or [])
                        images = q.get('images') or []
                    else:
                        # 兼容旧包格式（原始字段）
                        q_type = q.get('q_type') or q.get('题型') or '未知'
                        content = q.get('content') or q.get('题干') or ''
                        answer = q.get('answer') or q.get('答案') or ''
                        explanation = q.get('explanation') or q.get('解析') or ''
                        diff_val = q.get('difficulty') or q.get('难度') or 1
                        tags_str = tags_to_storage_str(q.get('tags') or q.get('标签') or [])
                        options_val = q.get('options') or q.get('选项') or '[]'
                        if isinstance(options_val, list):
                            options_json = json.dumps(options_val, ensure_ascii=False)
                        else:
                            options_json = str(options_val or '[]')

                        images = q.get('images') or []
                        if not images:
                            for img_rel in _normalize_image_paths(q.get('image_path')):
                                img_rel = str(img_rel).replace('\\', '/').lstrip('/')
                                images.append('images/' + img_rel)

                    # 3. 处理图片（统一为 JSON 列表字符串存储）
                    saved_paths = []
                    if images and isinstance(images, list):
                        for arcname in images:
                            arcname = str(arcname).replace('\\', '/').lstrip('/')
                            if arcname not in zf.namelist():
                                continue
                            ext = os.path.splitext(arcname)[1] or '.png'
                            unique_filename = f"{int(datetime.datetime.now().timestamp())}_{uuid.uuid4().hex}{ext}"
                            image_save_dir = os.path.join(upload_folder, 'question_images')
                            os.makedirs(image_save_dir, exist_ok=True)
                            image_save_path = os.path.join(image_save_dir, unique_filename)

                            with zf.open(arcname) as source, open(image_save_path, 'wb') as target:
                                target.write(source.read())
                            saved_paths.append(f'question_images/{unique_filename}')

                    image_path_val = json.dumps(saved_paths, ensure_ascii=False) if saved_paths else '[]'

                    # 4. 插入题目数据 (忽略原始ID，统一写入 PQF 同名列)
                    portable = internal_question_to_portable(
                        q_id=None,
                        q_type=q_type,
                        content=content,
                        options=options_json,
                        answer=answer,
                        explanation=explanation,
                        difficulty=diff_val,
                        tags=tags_str,
                    )

                    cols = ['subject_id', 'type', 'content', 'options', 'answer', 'analysis']
                    vals = [
                        subject_id,
                        portable.get('type') or 'essay',
                        portable.get('content') or '',
                        json.dumps(portable.get('options') or [], ensure_ascii=False),
                        json.dumps(
                            portable.get('answer') if portable.get('answer') is not None else [],
                            ensure_ascii=False,
                        ),
                        portable.get('analysis') or '',
                    ]

                    if has_difficulty:
                        cols.append('difficulty')
                        vals.append(int(portable.get('difficulty') or 1))
                    if has_tags:
                        cols.append('tags')
                        vals.append(json.dumps(portable.get('tags') or [], ensure_ascii=False))
                    if has_image_path:
                        cols.append('image_path')
                        vals.append(image_path_val)
                    if has_created_by:
                        cols.append('created_by')
                        vals.append(session.get('user_id') or q.get('created_by'))

                    if has_updated_at:
                        sql = f"INSERT INTO questions ({', '.join(cols)}, updated_at) VALUES ({', '.join(['?'] * len(vals))}, CURRENT_TIMESTAMP)"
                    else:
                        sql = f"INSERT INTO questions ({', '.join(cols)}) VALUES ({', '.join(['?'] * len(vals))})"

                    conn.execute(sql, vals)
                    imported_count += 1
                except Exception as e:
                    errors.append(f"导入题目ID {q.get('id', 'N/A')} 时出错: {str(e)}")

        conn.commit()
        try:
            bump_questions_version()
        except Exception:
            pass
        try:
            bump_subjects_version()
        except Exception:
            pass

        message = f'成功导入 {imported_count} 道题。'
        if errors:
            message += f' 遇到 {len(errors)} 个问题。'
        
        return jsonify({
            'status': 'success' if not errors else 'warning',
            'message': message,
            'imported_count': imported_count,
            'errors': errors
        })

    except zipfile.BadZipFile:
        return jsonify({'status': 'error', 'message': '文件不是一个有效的ZIP压缩包'}), 400
    except Exception as e:
        conn.rollback() # 如果发生意外错误，回滚事务
        return jsonify({'status': 'error', 'message': f'处理文件时发生未知错误: {str(e)}'}), 500
