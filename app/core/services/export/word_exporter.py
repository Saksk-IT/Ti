# -*- coding: utf-8 -*-
"""Word 导出生成器（python-docx，视觉优化版）"""
from __future__ import annotations

import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.core.utils.portable_question_format import (
    portable_type_to_q_type,
    normalize_portable_type,
)

from .base import ExportRequest, ExportResult, build_filename
from .formatter import ContentSegment, split_content, format_fill_blanks

# 颜色常量
_CLR_TITLE = RGBColor(0x1A, 0x1A, 0x2E)       # 近黑
_CLR_GROUP_TITLE = RGBColor(0x2B, 0x6C, 0xB0)  # 蓝色
_CLR_BODY = RGBColor(0x22, 0x22, 0x22)          # 题干正文
_CLR_OPTION = RGBColor(0x44, 0x44, 0x44)        # 选项深灰
_CLR_ANSWER_LABEL = RGBColor(0x1B, 0x8C, 0x5A)  # 答案绿
_CLR_ANSWER_TEXT = RGBColor(0x1B, 0x8C, 0x5A)
_CLR_ANALYSIS = RGBColor(0x77, 0x77, 0x77)      # 解析灰
_CLR_CODE_BG = "F5F5F5"
_CLR_CODE_BAR = "2B6CB0"

# 中文题型分组顺序
_TYPE_ORDER = ["single_choice", "multi_choice", "boolean", "fill", "essay"]
_TYPE_LABELS = {
    "single_choice": "选择题",
    "multi_choice": "多选题",
    "boolean": "判断题",
    "fill": "填空题",
    "essay": "简答题",
}
_CN_NUMBERS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]


def generate_word(req: ExportRequest, questions: list[dict[str, Any]]) -> ExportResult:
    """生成 Word 文档并返回 ExportResult。"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    # 文档标题
    _add_title(doc, req.subject_name)

    # 按题型分组
    grouped = _group_by_type(questions)

    group_idx = 0
    for p_type in _TYPE_ORDER:
        items = grouped.get(p_type, [])
        if not items:
            continue
        label = _TYPE_LABELS.get(p_type, p_type)
        cn = _CN_NUMBERS[group_idx] if group_idx < len(_CN_NUMBERS) else str(group_idx + 1)
        _add_group_heading(doc, f"{cn}、{label}（共 {len(items)} 题）")
        for i, q in enumerate(items, 1):
            _add_question(doc, q, i, p_type, req.include_answer)
            if i < len(items):
                _add_separator(doc)
        group_idx += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = build_filename(req, "docx")

    return ExportResult(
        buffer=buf,
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _group_by_type(questions: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    grouped: dict[str, list[dict[str, Any]]] = {}
    for q in questions:
        pt = normalize_portable_type(q.get("type", ""))
        if pt not in _TYPE_ORDER:
            pt = "essay"
        grouped.setdefault(pt, []).append(q)
    return grouped


def _add_title(doc: Document, subject_name: str) -> None:
    date_str = datetime.datetime.now().strftime("%Y年%m月%d日")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{subject_name}")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = _CLR_TITLE
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"导出日期：{date_str}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = _CLR_ANALYSIS
    run2.font.name = "微软雅黑"
    run2.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _add_group_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = _CLR_GROUP_TITLE
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _add_separator(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def _add_question(
    doc: Document,
    q: dict[str, Any],
    num: int,
    p_type: str,
    include_answer: bool,
) -> None:
    content = str(q.get("content") or "")
    options = q.get("options") or []
    answer_raw = q.get("answer")
    analysis = str(q.get("analysis") or "")

    if p_type == "fill":
        content = format_fill_blanks(content)

    # 题号 + 题干
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    num_run = p.add_run(f"{num}. ")
    num_run.font.bold = True
    num_run.font.size = Pt(11)
    num_run.font.color.rgb = _CLR_BODY
    num_run.font.name = "微软雅黑"
    num_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 题干内容（含代码块检测）
    segments = split_content(content)
    _render_segments_to_paragraph(doc, p, segments)

    # 选项
    if options and isinstance(options, list):
        for i, opt in enumerate(options):
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Cm(1.2)
            op.paragraph_format.space_before = Pt(1)
            op.paragraph_format.space_after = Pt(1)
            letter = chr(ord("A") + i) if i < 26 else str(i + 1)
            opt_text = str(opt).strip()
            # 去掉已有的 "A. " 前缀
            if len(opt_text) > 2 and opt_text[0].isalpha() and opt_text[1] in (".", "．", "、"):
                opt_text = opt_text[2:].strip()
            run = op.add_run(f"{letter}. {opt_text}")
            run.font.size = Pt(10.5)
            run.font.color.rgb = _CLR_OPTION
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    if not include_answer:
        return

    # 答案
    answer_str = _format_answer(answer_raw, p_type)
    if answer_str:
        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(4)
        label = ap.add_run("答案：")
        label.font.size = Pt(10)
        label.font.bold = True
        label.font.color.rgb = _CLR_ANSWER_LABEL
        label.font.name = "微软雅黑"
        label.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        ans_run = ap.add_run(answer_str)
        ans_run.font.size = Pt(10)
        ans_run.font.color.rgb = _CLR_ANSWER_TEXT
        ans_run.font.name = "微软雅黑"
        ans_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 解析
    if analysis:
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(2)
        elabel = ep.add_run("解析：")
        elabel.font.size = Pt(10)
        elabel.font.bold = True
        elabel.font.color.rgb = _CLR_ANALYSIS
        elabel.font.name = "微软雅黑"
        elabel.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        analysis_segs = split_content(analysis)
        _render_segments_to_paragraph(doc, ep, analysis_segs, font_size=Pt(10), color=_CLR_ANALYSIS)


def _render_segments_to_paragraph(
    doc: Document,
    para,
    segments: list[ContentSegment],
    *,
    font_size: Pt | None = None,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    """将 ContentSegment 列表渲染到段落/文档中。代码块单独用表格。"""
    size = font_size or Pt(11)
    clr = color or _CLR_BODY

    for seg in segments:
        if seg.kind == "code_block":
            _add_code_block_table(doc, seg.text, seg.language)
        elif seg.kind == "inline_code":
            run = para.add_run(seg.text)
            run.font.size = Pt(9.5)
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = para.add_run(seg.text)
            run.font.size = size
            run.font.color.rgb = clr
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if italic:
                run.font.italic = True


def _add_code_block_table(doc: Document, code: str, language: str) -> None:
    """用单行单列表格模拟代码块：浅灰背景 + 左侧蓝色竖线。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.text = ""

    # 浅灰背景
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): _CLR_CODE_BG,
    })
    tc_pr.append(shading)

    # 左侧蓝色边框
    tc_borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    left_border = tc_borders.makeelement(qn("w:left"), {
        qn("w:val"): "single",
        qn("w:sz"): "18",
        qn("w:space"): "0",
        qn("w:color"): _CLR_CODE_BAR,
    })
    tc_borders.append(left_border)
    tc_pr.append(tc_borders)

    p = cell.paragraphs[0]
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _format_answer(answer_raw: Any, p_type: str) -> str:
    """将 portable answer 格式化为可读字符串。"""
    if answer_raw is None:
        return ""
    if isinstance(answer_raw, list):
        if not answer_raw:
            return ""
        if p_type in ("single_choice", "multi_choice"):
            letters = []
            for v in answer_raw:
                try:
                    idx = int(v)
                    if 0 <= idx < 26:
                        letters.append(chr(ord("A") + idx))
                except (ValueError, TypeError):
                    pass
            return "".join(sorted(set(letters))) if letters else str(answer_raw)
        if p_type == "boolean":
            v = answer_raw[0] if answer_raw else None
            if v is True:
                return "正确"
            if v is False:
                return "错误"
            return str(v)
        if p_type == "fill":
            parts = []
            for i, group in enumerate(answer_raw):
                if isinstance(group, list):
                    parts.append(f"空{i+1}: {' / '.join(str(x) for x in group)}")
                else:
                    parts.append(f"空{i+1}: {group}")
            return "；".join(parts)
        # essay
        return "\n".join(str(x) for x in answer_raw)
    return str(answer_raw)
